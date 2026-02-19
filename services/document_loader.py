"""
services/document_loader.py
Load documents (PDF, DOCX, TXT) and split them into chunks.
All exceptions are wrapped in PDFLoadError for a clean boundary.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import List, Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter

import config.settings as cfg
from core.exceptions import PDFLoadError

logger = logging.getLogger(__name__)


def _splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=cfg.CHUNK_SIZE,
        chunk_overlap=cfg.CHUNK_OVERLAP,
    )


# ── PDF ───────────────────────────────────────────────────────────────────────

def load_pdf_chunks(path: str | Path) -> List[str]:
    """
    Extract text from a PDF and return chunks.
    Falls back to pytesseract OCR when no selectable text is found.
    Raises PDFLoadError on unrecoverable failure.
    """
    import fitz  # PyMuPDF

    path = Path(path)
    if not path.exists():
        raise PDFLoadError(f"PDF not found: {path}")

    try:
        doc = fitz.open(str(path))
        text = "".join(page.get_text() + "\n" for page in doc)
        doc.close()
    except Exception as exc:
        raise PDFLoadError(f"Cannot open PDF '{path}': {exc}") from exc

    if not text.strip():
        logger.warning("No selectable text in PDF – attempting OCR…")
        text = _ocr_pdf(path)

    if not text.strip():
        raise PDFLoadError(f"No text could be extracted from '{path}' (text + OCR both failed).")

    chunks = _splitter().split_text(text)
    if not chunks:
        raise PDFLoadError(f"Chunking produced 0 chunks for '{path}'.")

    logger.info("PDF '%s' → %d chunks", path.name, len(chunks))
    return chunks


def _ocr_pdf(path: Path) -> str:
    """Attempt OCR via pytesseract. Returns empty string on failure."""
    try:
        import fitz
        import pytesseract
        from PIL import Image

        doc = fitz.open(str(path))
        pages_text: List[str] = []
        for page in doc:
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            pages_text.append(pytesseract.image_to_string(img))
        doc.close()
        return "\n".join(pages_text)
    except Exception as exc:
        logger.error("OCR failed: %s", exc)
        return ""


# ── DOCX ──────────────────────────────────────────────────────────────────────

def load_docx_chunks(path: str | Path) -> List[str]:
    """
    Extract paragraph text from a DOCX and return chunks.
    The original Document object is returned separately via load_docx_document()
    when structure-preserving export is needed.
    """
    from docx import Document

    path = Path(path)
    if not path.exists():
        raise PDFLoadError(f"DOCX not found: {path}")

    try:
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as exc:
        raise PDFLoadError(f"Cannot open DOCX '{path}': {exc}") from exc

    if not text.strip():
        raise PDFLoadError(f"DOCX '{path}' contains no extractable text.")

    chunks = _splitter().split_text(text)
    logger.info("DOCX '%s' → %d chunks", path.name, len(chunks))
    return chunks


def load_docx_document(path: str | Path):
    """Return the raw python-docx Document object (for structure-preserving export)."""
    from docx import Document
    return Document(str(path))


# ── TXT ───────────────────────────────────────────────────────────────────────

def load_txt_chunks(path: str | Path) -> List[str]:
    path = Path(path)
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:
        raise PDFLoadError(f"Cannot read TXT '{path}': {exc}") from exc

    chunks = _splitter().split_text(text)
    logger.info("TXT '%s' → %d chunks", path.name, len(chunks))
    return chunks


# ── Dispatcher ────────────────────────────────────────────────────────────────

def load_chunks(path: str | Path) -> List[str]:
    """Auto-dispatch by extension. Raises PDFLoadError on failure."""
    suffix = Path(path).suffix.lower()
    dispatch = {
        ".pdf":  load_pdf_chunks,
        ".docx": load_docx_chunks,
        ".doc":  load_docx_chunks,
        ".txt":  load_txt_chunks,
    }
    loader = dispatch.get(suffix)
    if loader is None:
        raise PDFLoadError(f"Unsupported file type: '{suffix}'")
    return loader(path)
