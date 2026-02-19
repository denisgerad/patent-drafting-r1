"""
services/export_service.py
Convert text/markdown → DOCX or PDF BytesIO streams.
"""
from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def to_docx(
    markdown_text: str,
    base_doc=None,          # python-docx Document for structure-preservation
    preserve_structure: bool = False,
) -> Optional[io.BytesIO]:
    """Return a BytesIO DOCX stream or None on failure."""
    try:
        if preserve_structure and base_doc:
            return _docx_preserve_structure(markdown_text, base_doc)
        return _docx_simple(markdown_text)
    except Exception as exc:
        logger.error("DOCX export failed: %s", exc)
        return None


def _docx_simple(text: str) -> io.BytesIO:
    from docx import Document

    doc = Document()
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith("### "):
            doc.add_heading(para[4:].strip(), level=3)
        elif para.startswith("## "):
            doc.add_heading(para[3:].strip(), level=2)
        elif para.startswith("# "):
            doc.add_heading(para[2:].strip(), level=1)
        else:
            doc.add_paragraph(para)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _docx_preserve_structure(updated_text: str, base_doc) -> io.BytesIO:
    """
    Clone base_doc structure and replace text where matched.
    Images, tables, and formatting are preserved verbatim.
    """
    from docx import Document
    from docx.shared import Inches

    new_doc = Document()
    updated_lines = [l.strip() for l in updated_text.split("\n") if l.strip()]
    used: set[int] = set()

    def _has_image(run) -> bool:
        try:
            xml = run._element.xml
            return "graphic" in xml or "blip" in xml
        except Exception:
            return False

    def _copy_image(src_run, tgt_para, b_doc, n_doc) -> bool:
        try:
            drawings = src_run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
            for shape in drawings:
                blips = shape.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                )
                for blip in blips:
                    rId = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if rId:
                        img_bytes = b_doc.part.related_parts[rId].blob
                        extent = shape.find(
                            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent"
                        )
                        if extent is not None:
                            w = Inches(int(extent.get("cx")) / 914400)
                            h = Inches(int(extent.get("cy")) / 914400)
                            tgt_para.add_run().add_picture(io.BytesIO(img_bytes), width=w, height=h)
                        else:
                            tgt_para.add_run().add_picture(io.BytesIO(img_bytes))
                        return True
        except Exception:
            pass
        return False

    def _find_replacement(orig: str) -> Optional[str]:
        if not orig:
            return None
        for i, line in enumerate(updated_lines):
            if i in used:
                continue
            if line == orig or (len(orig) > 20 and line.startswith(orig[:30])):
                used.add(i)
                return line
        return None

    for element in base_doc.element.body:
        if element.tag.endswith("p"):
            orig_para = next(
                (p for p in base_doc.paragraphs if p._element is element), None
            )
            if orig_para is None:
                continue
            orig_text = orig_para.text.strip()
            replacement = _find_replacement(orig_text)
            new_para = new_doc.add_paragraph()
            new_para.style = orig_para.style

            if orig_para.runs:
                new_para.clear()
                text_placed = False
                for run in orig_para.runs:
                    if _has_image(run):
                        _copy_image(run, new_para, base_doc, new_doc)
                    else:
                        txt = (replacement if not text_placed else run.text) or run.text
                        if txt:
                            nr = new_para.add_run(txt)
                            nr.bold = run.bold
                            nr.italic = run.italic
                            nr.underline = run.underline
                            if run.font.size:
                                nr.font.size = run.font.size
                            if run.font.name:
                                nr.font.name = run.font.name
                            text_placed = True
            else:
                new_para.text = replacement or orig_text

        elif element.tag.endswith("tbl"):
            for table in base_doc.tables:
                if table._element is not element:
                    continue
                new_tbl = new_doc.add_table(rows=0, cols=len(table.columns))
                new_tbl.style = table.style
                for row in table.rows:
                    new_row = new_tbl.add_row()
                    for idx, cell in enumerate(row.cells):
                        nc = new_row.cells[idx]
                        nc.text = ""
                        for cp in cell.paragraphs:
                            cp2 = nc.add_paragraph()
                            cp2.style = cp.style
                            for run in cp.runs:
                                if _has_image(run):
                                    _copy_image(run, cp2, base_doc, new_doc)
                                else:
                                    r2 = cp2.add_run(run.text)
                                    r2.bold = run.bold
                                    r2.italic = run.italic
                break

    # Append unmatched updated lines
    for i, line in enumerate(updated_lines):
        if i not in used:
            new_doc.add_paragraph(line)

    bio = io.BytesIO()
    new_doc.save(bio)
    bio.seek(0)
    return bio


# ── PDF ───────────────────────────────────────────────────────────────────────

def to_pdf(text: str) -> Optional[io.BytesIO]:
    """Return a BytesIO PDF stream or None on failure."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas as rl_canvas

        bio = io.BytesIO()
        c = rl_canvas.Canvas(bio, pagesize=letter)
        w, h = letter
        margin, y, lh = 72, h - 72, 14

        for raw in text.splitlines():
            line = raw.strip()
            if not line:
                y -= lh
                continue
            text_out = line.lstrip("# ").upper() if line.startswith("# ") else line
            remaining = text_out
            while remaining:
                chunk = remaining[:90]
                c.drawString(margin, y, chunk)
                remaining = remaining[len(chunk):]
                y -= lh
                if y < margin:
                    c.showPage()
                    y = h - margin

        c.save()
        bio.seek(0)
        return bio
    except Exception as exc:
        logger.error("PDF export failed: %s", exc)
        return None
